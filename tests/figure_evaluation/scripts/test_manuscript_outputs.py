from __future__ import annotations

import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]


def _text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _docx_text(doc: Document) -> str:
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def test_manuscript_tex_and_pdf_outputs_exist() -> None:
    tex = ROOT / "manuscript" / "main.tex"
    pdf = ROOT / "manuscript" / "main.pdf"
    assert tex.exists()
    assert pdf.exists()
    assert pdf.stat().st_size > 4_000_000
    body = _text(tex)
    assert r"\documentclass[life,article,submit,moreauthors]{Definitions/mdpi}" in body
    assert "95.19\\% top-1 accuracy" in body
    assert "80.23\\% mean binary accuracy" in body
    assert len(re.findall(r"\\begin\{figure\}", body)) == 18
    assert "fig16_reproducibility_map" not in body
    assert "https://github.com/Mateo197802/titan-v45-ecg-multihead" in body
    assert "Architecture / parameter reporting" in body
    assert "Outputs and endpoint role" in body
    assert "Mastoi et al. (2022), Life" in body
    assert "This study & 12-lead modular specialist branches" not in body
    assert "Present specialist evidence package reported in this manuscript" in body
    assert "The present study is intentionally excluded from this state-of-the-art table" in body
    assert "Key challenges in multi-source ECG-AI reporting" in body
    assert "Primary6 row-normalized confusion matrix" in body
    assert "Primary6 confusion matrices" not in body
    assert "This work has five main limitations" not in body
    assert "Study Scope and Future Work" not in body
    assert "Public Repository" not in body
    assert "submitted to Life" not in body
    assert "life1010000" not in body
    assert "Interpretation." not in body
    assert len(re.findall(r"\\begin\{equation\}", body)) >= 12
    assert "MC dropout" in body or "MC-dropout" in body
    assert "Grad-CAM" in body
    assert "Optuna" in body
    assert len(re.findall(r"\\bibitem\{", body)) >= 45
    cls = _text(ROOT / "manuscript" / "Definitions" / "mdpi.cls")
    assert "submitted to {\\em \\journalname}" not in cls
    assert "Submitted to {\\em\\journalname}" not in cls
    assert "\\url{https://doi.org/\\@doinum}" not in cls
    assert "No DOI assigned" in cls


def test_manuscript_text_avoids_forbidden_claims() -> None:
    combined = "\n".join(
        _text(path)
        for path in [
            ROOT / "manuscript" / "main.tex",
            ROOT / "manuscript" / "main.md",
        ]
    )
    forbidden = [
        "TITAN V4",
        "teacher model",
        "ESP32",
        "accepted by decision",
        "gate passed",
        "reached the gate",
    ]
    lower = combined.lower()
    assert not [term for term in forbidden if term.lower() in lower]


def test_docx_contains_expected_content_and_assets() -> None:
    docx_path = ROOT / "manuscript" / "mdpi_life_full_manuscript_google_docs_ready_sanitized.docx"
    assert docx_path.exists()
    doc = Document(docx_path)
    text = _docx_text(doc)
    assert "Modular CNN-Transformer Specialist Cascades" in text
    assert "95.19% accuracy" in text
    assert "80.23% mean binary accuracy" in text
    assert "79.35% macro-F1" in text
    assert "Supplementary Materials" in text
    assert "Grad-CAM" in text
    assert "Interpretation." not in text
    assert "Public Repository" not in text
    assert "https://github.com/Mateo197802/titan-v45-ecg-multihead" in text
    assert "Study Scope and Future Work" not in text
    assert "This work has five main limitations" not in text
    assert "Architecture / parameter reporting" in text
    assert "Outputs and endpoint role" in text
    assert len(doc.tables) >= 9
    assert len(doc.inline_shapes) == 18
    assert text.count("https://doi.org/") >= 45
